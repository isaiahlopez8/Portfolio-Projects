from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import datetime
import os

#Creates a new Document
document = Document()

#Gets the current date
current_date = datetime.now().strftime('%B %d, %Y')
current_year = datetime.now().strftime('%Y')
current_month = datetime.now().strftime('%B')

#====================================================
#1 Add first line of paragraph and centers it
paragraph = document.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

#Adds a run. A run is a section of text within a paragrpah
run = paragraph.add_run('Surveillance\n')

#Set the run to be Bold and or Italic
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Calibri'
#====================================================
#2-4 Adds a run, we seperate to bold the previous
run = paragraph.add_run(f'Address & Information')

run.bold = True
run.font.size = Pt(14)
run.font.name = 'Calibri'
#====================================================
#Bullet Points. Must add a paragraph first. Then alignit. Then create a list of bullets.
#Then use a for statement to create a new paragraph for each bullet point. Uses List Bullet to use bullet style.

paragraph = document.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

bullet_points = [' ',' ',' ',]

for point in bullet_points:
    document.add_paragraph(point, style='List Bullet')
#====================================================
#Saves the document.
def Save():
    
    base_path = 'BasePath'

    year_folder = os.path.join(base_path, current_year)
    month_folder = os.path.join(year_folder, current_month)


# Ensure the folder exists
    os.makedirs(month_folder, exist_ok=True)

    file_name = f'Surveillance Report for the Day {current_date}.docx'
    file_path = os.path.join(month_folder, file_name)

    document.save(file_path) 
    #Opens the file 
    os.startfile(file_path)

Save()

